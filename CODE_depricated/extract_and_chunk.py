import fitz
import os
import json
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document



def extract_text_from_pdf(pdf_path):

    pdf = fitz.open(pdf_path)
    
    full_text = ""
    #Parse per page
    for page_num in range(len(pdf)):
        page = pdf[page_num]
        page_text = page.get_text()
        full_text += f"\n---Page {page_num + 1}---\n"
        full_text += page_text

    pdf.close()
    
    return full_text

def extract_text_from_docx(docx_path, include_tables=True):
    doc = Document(docx_path)
    full_text = "" 
     #Parse per paragraph since pages dont exist in docx until rendered
    for para in doc.paragraphs:
         if para.text.strip():  # Only add non-empty paragraphs
             if para.style.name.startswith('Heading'):
                 full_text += f"\n---{para.text.strip()}---\n"
             else:
                 full_text += para.text.strip() + "\n"
                
    if include_tables:
        try:
            for i, table in enumerate(doc.tables):
                full_text += f"\n---Table {i + 1}---\n"
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    full_text += "|".join(row_text + "\n")
                full_text += "\n"
        except:
            pass
    
    return full_text

#Step 2: Chunk the text
#Going with smaller size since the focus is on more detailed q&a. 400 tokens with 100 overlap based on testing with phi3:mini
def chunk_text(text, chunk_size=400, chunk_overlap=100):
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    chunks = text_splitter.split_text(text)
    
    return chunks
    
if __name__ == "__main__":
    from embedding import get_embedding
    from db import save_embeddings_to_db
    
    pdf_path = "mytaq_dna_polymerase_product_manual.pdf"
    docx_path = "example.docx"
    
    # Extract text from PDF
    extracted_pdf_text = extract_text_from_pdf(pdf_path)
    print(f"Extracted {len(extracted_pdf_text)} characters from PDF.")
    
    # Extract text from DOCX
    extracted_docx_text = extract_text_from_docx(docx_path)
    print(f"Extracted {len(extracted_docx_text)} characters from DOCX.")
    
    # Chunk the extracted PDF text
    pdf_chunks = chunk_text(extracted_pdf_text)
    print(f"Total PDF chunks: {len(pdf_chunks)}")
    
    # Chunk the extracted DOCX text
    docx_chunks = chunk_text(extracted_docx_text)
    print(f"Total DOCX chunks: {len(docx_chunks)}")
    
    # Chunk the extracted text
    pdf_chunks = chunk_text(extracted_pdf_text)
    docx_chunks = chunk_text(extracted_docx_text)
    
    print(f"Total PDF chunks: {len(pdf_chunks)}")
    print(f"Total DOCX chunks: {len(docx_chunks)}")
    
    # Create embeddings separately with metadata
    pdf_embeddings = [get_embedding(chunk) for chunk in pdf_chunks]
    docx_embeddings = [get_embedding(chunk) for chunk in docx_chunks]
    
    # Save with source metadata
    save_embeddings_to_db(pdf_embeddings, pdf_chunks, source="pdf", file_path=pdf_path)
    save_embeddings_to_db(docx_embeddings, docx_chunks, source="docx", file_path=docx_path)
    